from __future__ import annotations

import math
import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import Iterable

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# -----------------------------------------------------------------------------
# Instrumento e regras de pontuação
# -----------------------------------------------------------------------------
# A ordem abaixo corresponde ao questionário atualmente usado pelo projeto.
# Itens protetivos são invertidos para que, em todos os fatores, 0 represente
# menor exposição e 100 maior exposição percebida.

QUESTIONARIO: dict[int, str] = {
    1: "MINHA CHEFIA OFERECE A MIM E A MEUS COLEGAS BOAS OPORTUNIDADES DE DESENVOLVIMENTO",
    2: "MEU TRABALHO É RECONHECIDO PELA GERÊNCIA OU SUPERIORES",
    3: "MEUS SUPERIORES ME TRATAM DE FORMA JUSTA",
    4: "MEUS SUPERIORES ME FAZEM SENTIR PARTE DE UM GRUPO DE TRABALHO",
    5: "MINHA CHEFIA É BOA NO PLANEJAMENTO DO TRABALHO",
    6: "RECEBO TODA INFORMAÇÃO QUE NECESSITO PARA FAZER BEM O MEU TRABALHO",
    7: "CONFIO NA INFORMAÇÃO QUE É TRANSMITIDA PARA MIM PELOS MEUS SUPERIORES",
    8: "ESTOU SATISFEITO COM O MEU TRABALHO DE FORMA GERAL",
    9: "NO MEU TRABALHO, OS CONFLITOS SÃO RESOLVIDOS DE FORMA JUSTA",
    10: "MINHA CHEFIA CONFIA QUE SUA EQUIPE TEM CAPACIDADE PARA FAZER UM TRABALHO BEM FEITO",
    11: "AS TAREFAS SÃO BEM DISTRIBUÍDAS ENTRE OS COLEGAS DE TRABALHO",
    12: "TENHO AJUDA E APOIO DO MEU SUPERIOR IMEDIATO",
    13: "MEU CHEFE NÃO DEIXA CLARO QUAIS MUDANÇAS VÃO ACONTECER NO MEU TRABALHO",
    14: "APÓS UM DIA DE TRABALHO EU ME SINTO MUITO CANSADO(A) PELO EXCESSO DE TRABALHO",
    15: "SINTO QUE O MEU TRABALHO EXIGE DE MIM MUITA ENERGIA E QUE ISTO AFETA NEGATIVAMENTE MINHA VIDA FORA DO TRABALHO",
    16: "MEU TRABALHO ME TOMA MUITO TEMPO E ISTO PREJUDICA MINHA VIDA PESSOAL",
    17: "MEU TRABALHO ME DEIXA SEM ENERGIA",
    18: "PRECISO TRABALHAR MUITO RAPIDAMENTE",
    19: "PRECISO ACELERAR MUITO MEU RITMO DE TRABALHO PARA CUMPRIR MINHAS METAS DENTRO DO PRAZO",
    20: "A QUANTIDADE DE TAREFAS É EXCESSIVA PARA O TEMPO DISPONÍVEL PARA REALIZÁ-LAS",
    21: "MEU TRABALHO EXIGE QUE TOME DECISÕES DIFÍCEIS, O QUE FAZ COM QUE EU GASTE MUITA ENERGIA",
    22: "SINTO QUE O MEU TRABALHO É IMPORTANTE",
    23: "ME CONSIDERO CAPAZ DE RESOLVER PROBLEMAS SE FIZER O ESFORÇO NECESSÁRIO",
    24: "MINHAS RESPONSABILIDADES ME FAZEM SENTIR QUE MEU TRABALHO É IMPORTANTE",
    25: "O QUE APRENDO NO MEU TRABALHO FAZ COM QUE EU ME DESENVOLVA",
    26: "O MEU TRABALHO TEM SIGNIFICADO PARA MIM",
    27: "MEU TRABALHO ME PERMITE MOSTRAR QUE EU TENHO INICIATIVA",
    28: "CONSIGO INFLUENCIAR NA MANEIRA COMO FAÇO MEU TRABALHO",
}

# Itens redigidos como exposição/risco: quanto maior a frequência, maior o escore.
ITENS_EXPOSSICAO = set(range(13, 22))

FATORES = [
    {
        "codigo": "REL",
        "nome": "Relações interpessoais e liderança",
        "itens": list(range(1, 14)),
        "descricao": "Percepção de apoio, justiça, reconhecimento, comunicação e atuação da liderança.",
    },
    {
        "codigo": "DEM",
        "nome": "Demandas e intensidade do trabalho",
        "itens": list(range(14, 22)),
        "descricao": "Cansaço, ritmo, volume, pressão temporal e exigências cognitivas do trabalho.",
    },
    {
        "codigo": "INF",
        "nome": "Influência e significado do trabalho",
        "itens": list(range(22, 29)),
        "descricao": "Autonomia, iniciativa, desenvolvimento, autoeficácia percebida e sentido atribuído ao trabalho.",
    },
]

ACTION_GUIDES = {
    "REL": {
        "evidence": "apoio da liderança, justiça percebida, comunicação, reconhecimento e resolução de conflitos",
        "actions": [
            "realizar devolutiva com lideranças e trabalhadores para qualificar os achados",
            "revisar rotinas de comunicação de mudanças, critérios de distribuição de tarefas e reconhecimento",
            "definir procedimento de tratamento de conflitos e canais de escuta seguros",
        ],
        "indicator": "percentual de ações concluídas; reincidência de conflitos; evolução do escore no próximo ciclo",
    },
    "DEM": {
        "evidence": "volume, ritmo, pressão por prazo, cansaço e demandas cognitivas percebidas",
        "actions": [
            "reavaliar dimensionamento, prioridades, prazos e pausas reais",
            "mapear gargalos de processo e atividades com maior sobrecarga",
            "negociar medidas de organização do trabalho antes de responsabilizar indivíduos",
        ],
        "indicator": "horas extras; backlog; absenteísmo; percepção de sobrecarga no próximo ciclo",
    },
    "INF": {
        "evidence": "autonomia, influência, desenvolvimento, iniciativa e sentido do trabalho",
        "actions": [
            "ampliar participação dos trabalhadores em decisões sobre o modo de executar o trabalho",
            "criar oportunidades de desenvolvimento e uso de competências",
            "reforçar clareza de propósito, responsabilidades e margem de autonomia",
        ],
        "indicator": "adesão às ações; oportunidades implantadas; evolução do escore no próximo ciclo",
    },
}

PRIORITY_BY_LABEL = {
    "Prioridade": "0-30 dias",
    "Atenção": "30-60 dias",
    "Monitorar": "60-90 dias",
    "Condição favorável": "Manter e revisar no próximo ciclo",
    "Sem dados": "Coletar/qualificar dados antes de decidir",
}

RESPOSTAS_BASE = {
    "NUNCA / QUASE NUNCA": 0.0,
    "NUNCA QUASE NUNCA": 0.0,
    "NUNCA": 0.0,
    "RARAMENTE": 25.0,
    "AS VEZES": 50.0,
    "FREQUENTEMENTE": 75.0,
    "SEMPRE": 100.0,
}

ESCALA_ORDEM = [
    "NUNCA / QUASE NUNCA",
    "RARAMENTE",
    "ÀS VEZES",
    "FREQUENTEMENTE",
    "SEMPRE",
]

ESCALA_CURTA = ["Nunca/QN", "Raramente", "Às vezes", "Freq.", "Sempre"]

BASE_COLS = {
    "CARIMBO_DE_DATA_HORA",
    "CARIMBO_DE_DATAHORA",
    "NOME_DA_EMPRESA",
    "EMPRESA_ID",
    "CICLO_ID",
    "AREA_TRABALHO",
    "GRUPO_ID",
    "GRUPO_ANALISE",
    "SETOR",
    "CARGO",
    "CARGO_ATUAL",
    "IDADE",
    "NOME_COMPLETO",
    "CPF",
    "SEXO",
    "EMAIL",
    "E_MAIL",
    "CODIGO_RESPONDENTE",
    "PONTUACAO",
}


@dataclass(frozen=True)
class ReportConfig:
    min_group_size: int = 7
    min_answer_fraction: float = 0.70
    show_detailed_distribution: bool = True
    report_version: str = "4.0-v2-supabase-ready"
    methodology_status: str = "Triagem e monitoramento psicossocial; instrumento proprietário em validação técnica"


@dataclass
class GeneratedReport:
    scope: str
    respondents: int
    docx: bytes
    filename: str


# V2: sobrescreve as constantes originais da V1 mantendo compatibilidade com o restante do módulo.
from questionnaire_v2 import (  # noqa: E402
    ACTION_GUIDES,
    CRITICAL_ITEMS,
    FATORES,
    ITENS_EXPOSSICAO,
    OPEN_QUESTIONS,
    QUESTIONARIO,
)


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip().upper()
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_header(value: object) -> str:
    text = normalize_text(value)
    text = re.sub(r"[^A-Z0-9]+", "_", text).strip("_")
    return text


QUESTIONARIO_NORM = {number: normalize_header(text) for number, text in QUESTIONARIO.items()}


def normalize_dataframe_columns(df: pd.DataFrame) -> pd.DataFrame:
    result = df.copy()
    result.columns = [normalize_header(c) for c in result.columns]
    return result


def _clean_display_text(value: object) -> str:
    """Limpa o rótulo exibido sem retirar caixa ou acentos."""
    if value is None or pd.isna(value):
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _group_display_labels(
    keys: pd.Series,
    labels: pd.Series,
    fallback_labels: pd.Series | None = None,
) -> dict[str, str]:
    """Escolhe um rótulo legível e estável para cada chave canônica."""
    fallback_labels = labels if fallback_labels is None else fallback_labels
    display_by_key: dict[str, str] = {}
    fallback_by_key: dict[str, str] = {}
    for key, label, fallback in zip(keys, labels, fallback_labels):
        if not key:
            continue
        fallback_by_key.setdefault(key, _clean_display_text(fallback) or key)
        cleaned_label = _clean_display_text(label)
        if cleaned_label:
            display_by_key.setdefault(key, cleaned_label)
    return {
        key: display_by_key.get(key) or fallback
        for key, fallback in fallback_by_key.items()
    }


def _find_question_column(df: pd.DataFrame, question_number: int) -> str | None:
    expected = QUESTIONARIO_NORM[question_number]
    for col in df.columns:
        if normalize_header(col) == expected:
            return col
    return None


def answer_to_frequency_score(value: object) -> float:
    normalized = normalize_text(value)
    if normalized in RESPOSTAS_BASE:
        return RESPOSTAS_BASE[normalized]
    try:
        numeric = float(str(value).replace(",", "."))
    except (TypeError, ValueError):
        return float("nan")
    if 1 <= numeric <= 5:
        return (numeric - 1) * 25
    if 0 <= numeric <= 100:
        return numeric
    return float("nan")


def answer_to_exposure_score(value: object, question_number: int) -> float:
    score = answer_to_frequency_score(value)
    if math.isnan(score):
        return score
    if question_number not in ITENS_EXPOSSICAO:
        return 100.0 - score
    return score


def prepare_scored_dataframe(df_raw: pd.DataFrame) -> pd.DataFrame:
    df = normalize_dataframe_columns(df_raw)
    scored = pd.DataFrame(index=df.index)
    for question_number in QUESTIONARIO:
        source_col = _find_question_column(df, question_number)
        target_col = f"Q{question_number:02d}"
        if source_col is None:
            scored[target_col] = np.nan
        else:
            scored[target_col] = df[source_col].map(
                lambda value, q=question_number: answer_to_exposure_score(value, q)
            )
    return scored


def classify_exposure(score: float | int | None) -> str:
    """Faixa descritiva interna, não uma classificação regulatória autônoma."""
    if score is None or pd.isna(score):
        return "Sem dados"
    score = float(score)
    if score < 25:
        return "Condição favorável"
    if score < 50:
        return "Monitorar"
    if score < 75:
        return "Atenção"
    return "Prioridade"


def exposure_color(label: str) -> tuple[str, str]:
    palette = {
        "Condição favorável": ("D9EAD3", "274E13"),
        "Monitorar": ("FFF2CC", "7F6000"),
        "Atenção": ("FCE5CD", "783F04"),
        "Prioridade": ("F4CCCC", "990000"),
        "Sem dados": ("E7E6E6", "595959"),
    }
    return palette.get(label, palette["Sem dados"])


def calculate_factor_results(
    df_raw: pd.DataFrame,
    config: ReportConfig = ReportConfig(),
) -> pd.DataFrame:
    scored = prepare_scored_dataframe(df_raw)
    rows: list[dict[str, object]] = []
    for factor in FATORES:
        cols = [f"Q{i:02d}" for i in factor["itens"]]
        required = max(1, math.ceil(len(cols) * config.min_answer_fraction))
        valid_counts = scored[cols].notna().sum(axis=1)
        respondent_scores = scored[cols].mean(axis=1, skipna=True).where(valid_counts >= required)
        valid_scores = respondent_scores.dropna().astype(float)
        score = valid_scores.mean() if not valid_scores.empty else np.nan
        std = valid_scores.std(ddof=0) if len(valid_scores) > 1 else np.nan
        median = valid_scores.median() if not valid_scores.empty else np.nan
        high_pct = float((valid_scores >= 75).mean() * 100) if not valid_scores.empty else np.nan
        valid_respondents = int(valid_scores.shape[0])
        rows.append(
            {
                "Codigo": factor["codigo"],
                "Fator": factor["nome"],
                "Descricao": factor["descricao"],
                "Escore": round(float(score), 1) if pd.notna(score) else np.nan,
                "Mediana": round(float(median), 1) if pd.notna(median) else np.nan,
                "Desvio_padrao": round(float(std), 1) if pd.notna(std) else np.nan,
                "Pct_alta_exposicao": round(float(high_pct), 1) if pd.notna(high_pct) else np.nan,
                "Faixa": classify_exposure(score),
                "Respondentes_validos": valid_respondents,
            }
        )
    return pd.DataFrame(rows)


def canonical_answer_label(value: object) -> str | None:
    """Normaliza resposta textual ou numérica para a escala de frequência."""
    normalized = normalize_text(value)
    aliases = {
        "NUNCA / QUASE NUNCA": "NUNCA / QUASE NUNCA",
        "NUNCA QUASE NUNCA": "NUNCA / QUASE NUNCA",
        "NUNCA": "NUNCA / QUASE NUNCA",
        "RARAMENTE": "RARAMENTE",
        "AS VEZES": "ÀS VEZES",
        "ÀS VEZES": "ÀS VEZES",
        "ÀS VEZES": "ÀS VEZES",
        "FREQUENTEMENTE": "FREQUENTEMENTE",
        "SEMPRE": "SEMPRE",
    }
    if normalized in aliases:
        return aliases[normalized]
    try:
        numeric = float(str(value).replace(",", "."))
    except (TypeError, ValueError):
        return None
    if 1 <= numeric <= 5 and float(numeric).is_integer():
        return ESCALA_ORDEM[int(numeric) - 1]
    if 0 <= numeric <= 100:
        idx = min(4, max(0, round(numeric / 25)))
        return ESCALA_ORDEM[idx]
    return None


def validate_questionnaire_columns(df_raw: pd.DataFrame) -> dict[str, object]:
    """Retorna cobertura de colunas do instrumento sem expor dados individuais."""
    normalized = normalize_dataframe_columns(df_raw)
    missing = [number for number in QUESTIONARIO if _find_question_column(normalized, number) is None]
    found = [number for number in QUESTIONARIO if number not in missing]
    return {
        "expected_questions": len(QUESTIONARIO),
        "found_questions": len(found),
        "missing_questions": missing,
        "coverage": round(len(found) / len(QUESTIONARIO), 3) if QUESTIONARIO else 0,
    }


def calculate_question_results(df_raw: pd.DataFrame) -> pd.DataFrame:
    normalized = normalize_dataframe_columns(df_raw)
    question_to_factor: dict[int, dict[str, object]] = {}
    for factor in FATORES:
        for item in factor["itens"]:
            question_to_factor[int(item)] = factor

    rows: list[dict[str, object]] = []
    for number, question in QUESTIONARIO.items():
        source_col = _find_question_column(normalized, number)
        if source_col is None:
            continue
        exposure_scores = normalized[source_col].map(lambda value, q=number: answer_to_exposure_score(value, q))
        frequency_scores = normalized[source_col].map(answer_to_frequency_score)
        valid_scores = exposure_scores.dropna().astype(float)
        valid = int(valid_scores.shape[0])
        factor = question_to_factor.get(number, {})
        high_pct = float((valid_scores >= 75).mean() * 100) if valid else np.nan
        row: dict[str, object] = {
            "Numero": number,
            "Dominio_codigo": factor.get("codigo", ""),
            "Dominio": factor.get("nome", ""),
            "Pergunta": question,
            "Item_critico": number in CRITICAL_ITEMS,
            "Escore_exposicao": round(float(valid_scores.mean()), 1) if valid else np.nan,
            "Mediana": round(float(valid_scores.median()), 1) if valid else np.nan,
            "Desvio_padrao": round(float(valid_scores.std(ddof=0)), 1) if valid > 1 else np.nan,
            "Pct_alta_exposicao": round(float(high_pct), 1) if valid else np.nan,
            "Faixa": classify_exposure(valid_scores.mean() if valid else np.nan),
            "Respostas_validas": valid,
        }
        canonical_answers = normalized[source_col].map(canonical_answer_label)
        total = int(canonical_answers.notna().sum())
        for full_label, short_label in zip(ESCALA_ORDEM, ESCALA_CURTA):
            count = int((canonical_answers == full_label).sum())
            pct = (count / total * 100) if total else 0.0
            row[short_label] = f"{count} ({pct:.1f}%)"
        # Mantém para auditoria e depuração, sem expor resposta individual.
        row["Media_frequencia"] = round(float(frequency_scores.mean()), 1) if valid else np.nan
        rows.append(row)
    return pd.DataFrame(rows)


def build_critical_alerts(question_results: pd.DataFrame) -> list[dict[str, object]]:
    """Gera alertas sentinela para revisão técnica; não é diagnóstico ou conclusão jurídica."""
    alerts: list[dict[str, object]] = []
    if question_results.empty or "Item_critico" not in question_results.columns:
        return alerts
    critical = question_results[question_results["Item_critico"] == True].copy()  # noqa: E712
    for _, row in critical.iterrows():
        pct_high = row.get("Pct_alta_exposicao", np.nan)
        score = row.get("Escore_exposicao", np.nan)
        if pd.isna(pct_high) and pd.isna(score):
            continue
        severity = None
        if pd.notna(pct_high) and float(pct_high) >= 20:
            severity = "crítico"
        elif pd.notna(pct_high) and float(pct_high) >= 10:
            severity = "alto"
        elif pd.notna(score) and float(score) >= 50:
            severity = "atenção"
        if severity:
            alerts.append(
                {
                    "severidade": severity,
                    "numero": int(row["Numero"]),
                    "pergunta": str(row["Pergunta"]),
                    "escore": None if pd.isna(score) else float(score),
                    "pct_alta_exposicao": None if pd.isna(pct_high) else float(pct_high),
                    "mensagem": "Sinal crítico para revisão técnica, acolhimento seguro e definição de medidas proporcionais.",
                }
            )
    return alerts

def top_attention_items(question_results: pd.DataFrame, limit: int = 5) -> pd.DataFrame:
    if question_results.empty:
        return question_results
    return (
        question_results.dropna(subset=["Escore_exposicao"])
        .sort_values(["Escore_exposicao", "Numero"], ascending=[False, True])
        .head(limit)
        .copy()
    )


def build_recommendations(factor_results: pd.DataFrame) -> list[str]:
    recommendations: list[str] = []
    for _, row in factor_results.iterrows():
        code = str(row["Codigo"])
        label = str(row["Faixa"])
        guide = ACTION_GUIDES.get(code, {})
        if label in {"Sem dados", "Condição favorável"}:
            if label == "Condição favorável":
                recommendations.append(
                    f"{row['Fator']}: preservar as práticas positivas identificadas, documentar o que está funcionando e acompanhar sua manutenção nos próximos ciclos."
                )
            continue
        prefix = "priorizar" if label == "Prioridade" else "planejar"
        actions = "; ".join(guide.get("actions", []))
        recommendations.append(f"{row['Fator']}: {prefix} ações para {actions}.")
    recommendations.append(
        "Integrar os achados com observação do trabalho real, entrevistas, participação dos trabalhadores e avaliação técnica; "
        "o questionário isolado não substitui o processo de gerenciamento de riscos ocupacionais."
    )
    return recommendations


def build_action_plan(factor_results: pd.DataFrame) -> pd.DataFrame:
    """Plano de ação inicial para ser validado em devolutiva técnica pós-entrega."""
    rows: list[dict[str, object]] = []
    priority_rank = {"Prioridade": 0, "Atenção": 1, "Monitorar": 2, "Condição favorável": 3, "Sem dados": 4}
    for _, row in factor_results.iterrows():
        code = str(row["Codigo"])
        label = str(row["Faixa"])
        guide = ACTION_GUIDES.get(code, {})
        actions = guide.get("actions", [])
        rows.append(
            {
                "Prioridade": priority_rank.get(label, 9),
                "Fator": row["Fator"],
                "Faixa": label,
                "Evidência a qualificar": guide.get("evidence", "percepções coletivas do escopo analisado"),
                "Ações sugeridas": " • ".join(actions) if actions else "qualificar dados e definir medidas com responsáveis técnicos",
                "Responsável sugerido": "Gestão + liderança da área + apoio técnico/psicólogo parceiro",
                "Prazo inicial": PRIORITY_BY_LABEL.get(label, "Definir na devolutiva"),
                "Indicador de acompanhamento": guide.get("indicator", "evolução do escore no próximo ciclo"),
            }
        )
    return pd.DataFrame(rows).sort_values(["Prioridade", "Fator"]).drop(columns=["Prioridade"])


# -----------------------------------------------------------------------------
# Helpers de DOCX
# -----------------------------------------------------------------------------

def _set_cell_text(
    cell,
    text: object,
    *,
    bold: bool = False,
    size: float = 9,
    color: str = "1F2937",
    align: str = "left",
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_bg(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_table_borders(table, color: str = "D1D5DB", size: str = "4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_picture_alt_text(inline_shape, *, title: str, description: str) -> None:
    """Adiciona título e texto alternativo ao gráfico para acessibilidade."""
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def _set_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size, color in [
        ("Title", 22, "16324F"),
        ("Heading 1", 15, "16324F"),
        ("Heading 2", 12, "26547C"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True


def _configure_portrait_section(section) -> None:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)


def _configure_landscape_section(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)


def _add_header_footer(doc: Document, company: str, cycle_label: str) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.text = ""
        run = p.add_run(f"{company} | {cycle_label}")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("6B7280")

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""
        fr = fp.add_run("Documento confidencial - resultados coletivos e descritivos")
        fr.font.name = "Arial"
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = RGBColor.from_string("6B7280")


def _add_cover(
    doc: Document,
    *,
    company: str,
    scope: str,
    cycle_label: str,
    respondents: int,
    generated_at: datetime,
    version: str,
    methodology_status: str,
) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(32)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO DE INDICADORES PSICOSSOCIAIS DO TRABALHO")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("16324F")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(company)
    run2.bold = True
    run2.font.name = "Arial"
    run2.font.size = Pt(17)
    run2.font.color.rgb = RGBColor.from_string("26547C")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(scope)
    r3.font.name = "Arial"
    r3.font.size = Pt(13)
    r3.font.color.rgb = RGBColor.from_string("4B5563")

    doc.add_paragraph("")
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(9.0)
    _set_table_borders(table, color="D9E2F3")
    info = [
        ("Ciclo", cycle_label),
        ("Respostas analisadas", respondents),
        ("Data de emissão", generated_at.strftime("%d/%m/%Y %H:%M")),
        ("Versão do relatório", version),
        ("Status metodológico", methodology_status),
    ]
    for row, (label, value) in zip(table.rows, info):
        _set_cell_bg(row.cells[0], "EAF2F8")
        _set_cell_text(row.cells[0], label, bold=True, size=9, color="16324F")
        _set_cell_text(row.cells[1], value, size=9.5)
        for cell in row.cells:
            _set_cell_margins(cell)

    doc.add_paragraph("")
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(14)
    rn = notice.add_run(
        "Uso organizacional. Este documento não realiza diagnóstico clínico individual e não substitui a avaliação técnica "
        "integrada prevista no gerenciamento de riscos ocupacionais."
    )
    rn.italic = True
    rn.font.name = "Arial"
    rn.font.size = Pt(8.5)
    rn.font.color.rgb = RGBColor.from_string("6B7280")
    doc.add_page_break()

def _add_scope_and_method(
    doc: Document,
    scope: str,
    suppressed_sectors: Iterable[str] | None = None,
    methodology_status: str = "Em validação técnica",
) -> None:
    doc.add_heading("1. Escopo e interpretação", level=1)
    p = doc.add_paragraph()
    p.add_run("Escopo analisado: ").bold = True
    p.add_run(scope)

    status = doc.add_paragraph()
    status.add_run("Status metodológico: ").bold = True
    status.add_run(methodology_status)

    paragraphs = [
        "Os escores são apresentados em escala de 0 a 100, orientada para exposição: quanto maior o valor, maior a percepção de condições que merecem investigação e prevenção.",
        "As faixas visuais são descritivas e operacionais. Elas não constituem, isoladamente, classificação legal de risco, diagnóstico de saúde, nexo causal, laudo pericial ou conclusão automática de conformidade normativa.",
        "A análise deve ser combinada com observação do trabalho real, entrevistas, participação dos trabalhadores, dados de SST/RH, avaliação ergonômica quando aplicável e julgamento técnico das medidas de prevenção.",
        "As respostas são tratadas de forma coletiva. Resultados de grupos abaixo do mínimo configurado são suprimidos para reduzir o risco de identificação direta ou indireta.",
        "A interpretação final deve ser conduzida por profissional habilitado/qualificado, considerando o contexto da organização e preservando a confidencialidade dos participantes.",
    ]
    for item in paragraphs:
        doc.add_paragraph(item, style="List Bullet")

    suppressed_count = len([s for s in (suppressed_sectors or []) if s])
    if suppressed_count:
        p2 = doc.add_paragraph()
        p2.add_run("Proteção de confidencialidade: ").bold = True
        noun = "recorte setorial foi suprimido" if suppressed_count == 1 else "recortes setoriais foram suprimidos"
        p2.add_run(
            f"{suppressed_count} {noun} por não atingir o mínimo de participantes. "
            "Os nomes desses grupos não são exibidos neste documento."
        )

def _add_factor_cards(doc: Document, factor_results: pd.DataFrame) -> None:
    doc.add_heading("2. Síntese dos fatores", level=1)
    doc.add_paragraph(
        "Além do escore médio, a tabela mostra mediana, dispersão e percentual de respondentes em alta exposição. "
        "Esses indicadores ajudam a diferenciar problemas homogêneos de situações concentradas em subgrupos."
    )
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(5.2), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.4), Cm(3.0)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["Fator", "Média", "Mediana", "Desvio", "% alta exp.", "Faixa"]
    for i, header in enumerate(headers):
        _set_cell_bg(table.rows[0].cells[i], "16324F")
        _set_cell_text(table.rows[0].cells[i], header, bold=True, size=8.3, color="FFFFFF", align="center")
    _repeat_table_header(table.rows[0])
    for _, row in factor_results.iterrows():
        cells = table.add_row().cells
        _set_cell_text(cells[0], row["Fator"], bold=True, size=8.2)
        values = [
            "-" if pd.isna(row.get("Escore")) else f"{float(row['Escore']):.1f}",
            "-" if pd.isna(row.get("Mediana")) else f"{float(row['Mediana']):.1f}",
            "-" if pd.isna(row.get("Desvio_padrao")) else f"{float(row['Desvio_padrao']):.1f}",
            "-" if pd.isna(row.get("Pct_alta_exposicao")) else f"{float(row['Pct_alta_exposicao']):.1f}%",
        ]
        for idx, value in enumerate(values, start=1):
            _set_cell_text(cells[idx], value, bold=idx == 1, size=8.5, align="center")
        bg, fg = exposure_color(str(row["Faixa"]))
        _set_cell_bg(cells[5], bg)
        _set_cell_text(cells[5], row["Faixa"], bold=True, size=8, color=fg, align="center")
        for cell in cells:
            _set_cell_margins(cell, top=55, bottom=55)
        _prevent_row_split(table.rows[-1])
    _set_table_borders(table)


def _build_factor_chart(factor_results: pd.DataFrame, title: str) -> BytesIO:
    chart_data = factor_results.copy()
    values = chart_data["Escore"].fillna(0).astype(float)
    labels = chart_data["Fator"].tolist()
    colors = []
    for label in chart_data["Faixa"]:
        bg, _ = exposure_color(str(label))
        colors.append(f"#{bg}")

    fig, ax = plt.subplots(figsize=(9.6, 4.8))
    bars = ax.barh(labels, values, color=colors, edgecolor="#8A94A6")
    ax.set_xlim(0, 100)
    ax.set_xlabel("Escore de exposição (0-100)")
    ax.set_title(title)
    ax.grid(axis="x", alpha=0.20)
    ax.invert_yaxis()
    for bar, value in zip(bars, values):
        ax.text(min(value + 1.5, 96), bar.get_y() + bar.get_height() / 2, f"{value:.1f}", va="center", fontsize=9)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _add_factor_chart(doc: Document, factor_results: pd.DataFrame, scope: str) -> None:
    chart = _build_factor_chart(factor_results, f"Indicadores por fator - {scope}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(chart, width=Inches(6.8))
    summary = "; ".join(
        f"{row['Fator']}: {row['Escore']} pontos, {row['Faixa']}"
        for _, row in factor_results.iterrows()
    )
    _set_picture_alt_text(
        shape,
        title=f"Indicadores por fator - {scope}",
        description=f"Gráfico de barras horizontais. {summary}.",
    )


def _add_attention_items(doc: Document, question_results: pd.DataFrame) -> None:
    doc.add_page_break()
    doc.add_heading("3. Itens com maior escore de exposição", level=1)
    top = top_attention_items(question_results, 6)
    if top.empty:
        doc.add_paragraph("Não foi possível calcular os itens por ausência de respostas válidas.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(1.0), Cm(10.0), Cm(2.2), Cm(3.2)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["Nº", "Pergunta", "Escore", "Faixa"]
    for i, header in enumerate(headers):
        _set_cell_bg(table.rows[0].cells[i], "26547C")
        _set_cell_text(table.rows[0].cells[i], header, bold=True, size=8.5, color="FFFFFF", align="center")
    _repeat_table_header(table.rows[0])
    for _, row in top.iterrows():
        cells = table.add_row().cells
        _set_cell_text(cells[0], int(row["Numero"]), bold=True, size=8.5, align="center")
        _set_cell_text(cells[1], str(row["Pergunta"]).capitalize(), size=8.5)
        _set_cell_text(cells[2], f"{float(row['Escore_exposicao']):.1f}", bold=True, size=9, align="center")
        bg, fg = exposure_color(str(row["Faixa"]))
        _set_cell_bg(cells[3], bg)
        _set_cell_text(cells[3], row["Faixa"], bold=True, size=8, color=fg, align="center")
        for cell in cells:
            _set_cell_margins(cell, top=60, bottom=60)
        _prevent_row_split(table.rows[-1])
    _set_table_borders(table)



def _add_critical_alerts(doc: Document, question_results: pd.DataFrame) -> None:
    alerts = build_critical_alerts(question_results)
    doc.add_heading("3.1 Alertas críticos para revisão técnica", level=2)
    doc.add_paragraph(
        "Os alertas abaixo são sinais sentinela. Eles não confirmam assédio, violência, nexo causal ou irregularidade por si só; "
        "indicam necessidade de acolhimento seguro, análise técnica e, quando aplicável, fluxo de apuração e medidas preventivas."
    )
    if not alerts:
        doc.add_paragraph("Não foram identificados alertas críticos pelos critérios automáticos desta triagem.")
        return
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(1.6), Cm(1.2), Cm(9.2), Cm(2.0), Cm(2.0)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["Severidade", "Item", "Sinal", "Escore", "% alta"]
    for i, header in enumerate(headers):
        _set_cell_bg(table.rows[0].cells[i], "7F1D1D")
        _set_cell_text(table.rows[0].cells[i], header, bold=True, size=8, color="FFFFFF", align="center")
    _repeat_table_header(table.rows[0])
    for alert in alerts:
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(alert["severidade"]).upper(), bold=True, size=7.8, align="center")
        _set_cell_text(cells[1], alert["numero"], bold=True, size=8, align="center")
        _set_cell_text(cells[2], str(alert["pergunta"]), size=7.8)
        _set_cell_text(cells[3], "-" if alert["escore"] is None else f"{alert['escore']:.1f}", size=8, align="center")
        _set_cell_text(cells[4], "-" if alert["pct_alta_exposicao"] is None else f"{alert['pct_alta_exposicao']:.1f}%", size=8, align="center")
        for cell in cells:
            _set_cell_margins(cell, top=45, bottom=45, start=30, end=30)
        _prevent_row_split(table.rows[-1])
    _set_table_borders(table, size="3")

def _add_recommendations(doc: Document, factor_results: pd.DataFrame) -> None:
    doc.add_heading("4. Recomendações e plano de ação inicial", level=1)
    for recommendation in build_recommendations(factor_results):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(recommendation)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_landscape_section(section)
    doc.add_heading("4.1 Plano de ação para validação na devolutiva", level=2)
    doc.add_paragraph(
        "A tabela abaixo é uma proposta inicial. Antes de implementação, recomenda-se validá-la com a empresa, trabalhadores, liderança e responsável técnico, considerando evidências adicionais e viabilidade operacional."
    )
    action_plan = build_action_plan(factor_results)
    cols = ["Fator", "Faixa", "Evidência a qualificar", "Ações sugeridas", "Responsável sugerido", "Prazo inicial", "Indicador de acompanhamento"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.4), Cm(2.1), Cm(3.4), Cm(5.4), Cm(3.5), Cm(2.6), Cm(3.8)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for i, header in enumerate(cols):
        _set_cell_bg(table.rows[0].cells[i], "16324F")
        _set_cell_text(table.rows[0].cells[i], header, bold=True, size=7.2, color="FFFFFF", align="center")
        _set_cell_margins(table.rows[0].cells[i], top=45, bottom=45, start=35, end=35)
    _repeat_table_header(table.rows[0])
    for _, row in action_plan.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            value = row.get(col, "")
            _set_cell_text(cells[i], value, size=6.8, align="center" if col in {"Faixa", "Prazo inicial"} else "left")
            _set_cell_margins(cells[i], top=40, bottom=40, start=30, end=30)
        bg, fg = exposure_color(str(row.get("Faixa", "")))
        _set_cell_bg(cells[1], bg)
        for run in cells[1].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(fg)
            run.bold = True
        _prevent_row_split(table.rows[-1])
    _set_table_borders(table, size="3")

def _add_distribution_appendix(doc: Document, question_results: pd.DataFrame) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_landscape_section(section)

    rows_per_page = 14
    chunks = [
        question_results.iloc[start : start + rows_per_page]
        for start in range(0, len(question_results), rows_per_page)
    ] or [question_results]

    for part_index, chunk in enumerate(chunks, start=1):
        if part_index > 1:
            doc.add_page_break()

        heading = "Anexo - distribuição das respostas"
        if len(chunks) > 1:
            heading += f" ({part_index}/{len(chunks)})"
        doc.add_heading(heading, level=1)
        doc.add_paragraph(
            "A distribuição ajuda a identificar heterogeneidade das percepções. "
            "O escore de exposição já considera o sentido positivo ou negativo de cada item."
        )

        cols = ["Numero", "Pergunta", "Escore_exposicao", "Faixa", *ESCALA_CURTA]
        table = doc.add_table(rows=1, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = [Cm(0.8), Cm(10.8), Cm(1.4), Cm(2.3), Cm(1.9), Cm(1.9), Cm(1.9), Cm(1.9), Cm(1.9)]
        for idx, width in enumerate(widths):
            table.columns[idx].width = width

        labels = ["Nº", "Pergunta", "Escore", "Faixa", *ESCALA_CURTA]
        for i, label in enumerate(labels):
            _set_cell_bg(table.rows[0].cells[i], "16324F")
            _set_cell_text(
                table.rows[0].cells[i],
                label,
                bold=True,
                size=7.8,
                color="FFFFFF",
                align="center",
            )
            _set_cell_margins(table.rows[0].cells[i], top=45, bottom=45, start=35, end=35)
        _repeat_table_header(table.rows[0])

        for _, row in chunk.iterrows():
            cells = table.add_row().cells
            values = [
                int(row["Numero"]),
                str(row["Pergunta"]).capitalize(),
                "-" if pd.isna(row["Escore_exposicao"]) else f"{float(row['Escore_exposicao']):.1f}",
                row["Faixa"],
                *[row.get(label, "") for label in ESCALA_CURTA],
            ]
            for i, value in enumerate(values):
                _set_cell_text(cells[i], value, size=7.8, align="left" if i == 1 else "center")
                _set_cell_margins(cells[i], top=45, bottom=45, start=35, end=35)
            bg, fg = exposure_color(str(row["Faixa"]))
            _set_cell_bg(cells[3], bg)
            for run in cells[3].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(fg)
                run.bold = True
            _prevent_row_split(table.rows[-1])
        _set_table_borders(table, size="3")

def _safe_filename(text: str) -> str:
    normalized = normalize_text(text)
    normalized = re.sub(r"[^A-Z0-9]+", "_", normalized).strip("_")
    return normalized[:80] or "RELATORIO"


def create_collective_report(
    df_scope: pd.DataFrame,
    *,
    company: str,
    scope: str,
    cycle_label: str,
    generated_at: datetime | None = None,
    suppressed_sectors: Iterable[str] | None = None,
    config: ReportConfig = ReportConfig(),
) -> GeneratedReport:
    if df_scope.empty:
        raise ValueError("Não há respostas no escopo informado.")
    generated_at = generated_at or datetime.now()
    factor_results = calculate_factor_results(df_scope, config)
    question_results = calculate_question_results(df_scope)

    doc = Document()
    _set_doc_styles(doc)
    _configure_portrait_section(doc.sections[0])
    _add_cover(
        doc,
        company=company,
        scope=scope,
        cycle_label=cycle_label,
        respondents=len(df_scope),
        generated_at=generated_at,
        version=config.report_version,
        methodology_status=config.methodology_status,
    )
    _add_scope_and_method(
        doc,
        scope,
        suppressed_sectors,
        methodology_status=config.methodology_status,
    )
    _add_factor_cards(doc, factor_results)
    _add_factor_chart(doc, factor_results, scope)
    _add_attention_items(doc, question_results)
    _add_critical_alerts(doc, question_results)
    _add_recommendations(doc, factor_results)
    if config.show_detailed_distribution:
        _add_distribution_appendix(doc, question_results)
    _add_header_footer(doc, company, cycle_label)

    output = BytesIO()
    doc.save(output)
    filename = (
        f"Relatorio_{_safe_filename(company)}_{_safe_filename(scope)}_"
        f"{generated_at.strftime('%Y%m%d')}.docx"
    )
    return GeneratedReport(scope=scope, respondents=len(df_scope), docx=output.getvalue(), filename=filename)


def generate_company_reports(
    df_company: pd.DataFrame,
    *,
    company: str,
    cycle_label: str,
    config: ReportConfig = ReportConfig(),
    generated_at: datetime | None = None,
) -> tuple[list[GeneratedReport], list[str]]:
    """Gera o relatório geral e recortes por grupo de análise canônico."""
    if df_company.empty:
        raise ValueError("Não existem respostas para a empresa/ciclo.")
    generated_at = generated_at or datetime.now()
    normalized = normalize_dataframe_columns(df_company)
    reports: list[GeneratedReport] = []
    suppressed: list[str] = []

    canonical_columns = {"GRUPO_ID", "GRUPO_ANALISE"}.issubset(normalized.columns)
    canonical_keys = (
        normalized["GRUPO_ID"].fillna("").map(_clean_display_text)
        if canonical_columns
        else pd.Series("", index=normalized.index, dtype="object")
    )

    if canonical_columns and canonical_keys.ne("").any():
        raw_ids = normalized["GRUPO_ID"].fillna("")
        group_keys = canonical_keys
        display_by_key = _group_display_labels(
            group_keys,
            normalized["GRUPO_ANALISE"].fillna(""),
            raw_ids,
        )
        scope_prefix = "Grupo de análise"
    elif "SETOR" in normalized.columns:
        raw_sectors = normalized["SETOR"].fillna("")
        group_keys = raw_sectors.map(normalize_text)
        display_by_key = _group_display_labels(group_keys, raw_sectors)
        scope_prefix = "Setor"
    else:
        group_keys = pd.Series("", index=normalized.index, dtype="object")
        display_by_key = {}
        scope_prefix = "Grupo de análise"

    counts = group_keys[group_keys != ""].value_counts()
    sort_key = lambda key: (normalize_text(display_by_key.get(key, key)), key)
    eligible = sorted(
        [key for key, count in counts.items() if int(count) >= config.min_group_size],
        key=sort_key,
    )
    suppressed_keys = sorted(
        [key for key, count in counts.items() if int(count) < config.min_group_size],
        key=sort_key,
    )
    suppressed = [display_by_key.get(key, key) for key in suppressed_keys]

    reports.append(
        create_collective_report(
            normalized,
            company=company,
            scope="Visão geral da empresa",
            cycle_label=cycle_label,
            generated_at=generated_at,
            suppressed_sectors=suppressed,
            config=config,
        )
    )

    for group_key in eligible:
        subset = normalized[group_keys == group_key].copy()
        display_label = display_by_key.get(group_key, group_key)
        reports.append(
            create_collective_report(
                subset,
                company=company,
                scope=f"{scope_prefix}: {display_label}",
                cycle_label=cycle_label,
                generated_at=generated_at,
                config=config,
            )
        )
    return reports, suppressed
