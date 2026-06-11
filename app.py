import streamlit as st
import pandas as pd
import numpy as np
import re
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import gspread
from google.oauth2.service_account import Credentials

st.set_page_config(page_title="Gerador de Relatórios Psicossocial", page_icon="📊", layout="wide")

def norm_text(x):
    s = str(x).strip().upper()
    s = s.replace('Á', 'A').replace('À', 'A').replace('Â', 'A').replace('Ã', 'A')
    s = s.replace('É', 'E').replace('Ê', 'E')
    s = s.replace('Í', 'I')
    s = s.replace('Ó', 'O').replace('Ô', 'O').replace('Õ', 'O')
    s = s.replace('Ú', 'U').replace('Ü', 'U')
    s = s.replace('Ç', 'C')
    s = re.sub(r'\s+', ' ', s)
    return s

MAP = {
    1: 1, 2: 2, 3: 3, 4: 4, 5: 5,
    '1': 1, '2': 2, '3': 3, '4': 4, '5': 5,
    'NUNCA / QUASE NUNCA': 1, 'NUNCA QUASE NUNCA': 1, 'NUNCA': 1,
    'RARAMENTE': 2, 'ÀS VEZES': 3, 'AS VEZES': 3, 'S VEZES': 3,
    'FREQUENTEMENTE': 4, 'SEMPRE': 5
}

ESCALA_TEXTOS = ['NUNCA / QUASE NUNCA', 'RARAMENTE', 'ÀS VEZES', 'FREQUENTEMENTE', 'SEMPRE']
ESCALA_LABELS = ['Nunca/Q.Nunca', 'Raramente', 'Às vezes', 'Frequentemente', 'Sempre']

BLOCOS = [
    {'nome': 'Sobrecarga de trabalho', 'inicio': 1, 'fim': 13},
    {'nome': 'Conflito interpessoal e falta de apoio', 'inicio': 14, 'fim': 20},
    {'nome': 'Relações físicas e emocionais ao trabalho', 'inicio': 21, 'fim': 28}
]

QUESTIONARIO = {
    1: 'MINHA CHEFIA OFERECE A MIM E A MEUS COLEGAS BOAS OPORTUNIDADES DE DESENVOLVIMENTO',
    2: 'MEU TRABALHO É RECONHECIDO PELA GERÊNCIA OU SUPERIORES',
    3: 'MEUS SUPERIORES ME TRATAM DE FORMA JUSTA',
    4: 'MEUS SUPERIORES ME FAZEM SENTIR PARTE DE UM GRUPO DE TRABALHO',
    5: 'MINHA CHEFIA É BOA NO PLANEJAMENTO DO TRABALHO',
    6: 'RECEBO TODA INFORMAÇÃO QUE NECESSITO PARA FAZER BEM O MEU TRABALHO',
    7: 'CONFIO NA INFORMAÇÃO QUE É TRANSMITIDA PARA MIM PELOS MEUS SUPERIORES',
    8: 'ESTOU SATISFEITO COM O MEU TRABALHO DE FORMA GERAL',
    9: 'NO MEU TRABALHO, OS CONFLITOS SÃO RESOLVIDOS DE FORMA JUSTA',
    10: 'MINHA CHEFIA CONFIA QUE SUA EQUIPE TEM CAPACIDADE PARA FAZER UM TRABALHO BEM FEITO',
    11: 'AS TAREFAS SÃO BEM DISTRIBUÍDAS ENTRE OS COLEGAS DE TRABALHO',
    12: 'TENHO AJUDA E APOIO DO MEU SUPERIOR IMEDIATO',
    13: 'MEU CHEFE NÃO DEIXA CLARO QUAIS MUDANÇAS VÃO ACONTECER NO MEU TRABALHO',
    14: 'APÓS UM DIA DE TRABALHO EU ME SINTO MUITO CANSADO(A) PELO EXCESSO DE TRABALHO',
    15: 'SINTO QUE O MEU TRABALHO EXIGE DE MIM MUITA ENERGIA E QUE ISTO AFETA NEGATIVAMENTE MINHA VIDA FORA DO TRABALHO',
    16: 'MEU TRABALHO ME TOMA MUITO TEMPO E ISTO PREJUDICA MINHA VIDA PESSOAL',
    17: 'MEU TRABALHO ME DEIXA SEM ENERGIA',
    18: 'PRECISO TRABALHAR MUITO RAPIDAMENTE',
    19: 'PRECISO ACELERAR MUITO MEU RITMO DE TRABALHO PARA CUMPRIR MINHAS METAS DENTRO DO PRAZO',
    20: 'A QUANTIDADE DE TAREFAS É EXCESSIVA PARA O TEMPO DISPONÍVEL PARA REALIZÁ-LAS',
    21: 'MEU TRABALHO EXIGE QUE TOME DECISÕES DIFÍCEIS, O QUE FAZ COM QUE EU GASTE MUITA ENERGIA',
    22: 'SINTO QUE O MEU TRABALHO É IMPORTANTE',
    23: 'ME CONSIDERO CAPAZ DE RESOLVER PROBLEMAS SE FIZER O ESFORÇO NECESSÁRIO',
    24: 'MINHAS RESPONSABILIDADES ME FAZEM SENTIR QUE MEU TRABALHO É IMPORTANTE',
    25: 'O QUE APRENDO NO MEU TRABALHO FAZ COM QUE EU ME DESENVOLVA',
    26: 'O MEU TRABALHO TEM SIGNIFICADO PARA MIM',
    27: 'MEU TRABALHO ME PERMITE MOSTRAR QUE EU TENHO INICIATIVA',
    28: 'CONSIGO INFLUENCIAR NA MANEIRA COMO FAÇO MEU TRABALHO'
}

NEGATIVAS = {13, 14, 15, 16, 17, 18, 19, 20, 21}
BASE_COLS = {'CARIMBO DE DATA/HORA', 'CARIMBO DE DATAHORA', 'NOME DA EMPRESA', 'SETOR', 'CARGO', 'CARGO ATUAL', 'IDADE', 'NOME COMPLETO', 'CPF', 'SEXO', 'PONTUAÇÃO', 'PONTUACAO'}

QUESTIONARIO_NORM = {k: norm_text(v) for k, v in QUESTIONARIO.items()}

def convert_answer_to_num(x):
    s = norm_text(x)
    if s in MAP and not pd.isna(MAP[s]):
        return MAP[s]
    try:
        n = float(str(x).replace(',', '.'))
        if 1 <= n <= 5:
            return n
    except:
        pass
    return np.nan

def classify_by_media(media):
    if pd.isna(media):
        return 'Sem dados'
    if media <= 2.5:
        return 'Risco Baixo'
    if media <= 3.5:
        return 'Risco Médio'
    return 'Risco Alto'

def color_by_classificacao(clf):
    if clf == 'Risco Baixo':
        return ('70AD47', 'FFFFFF')
    if clf == 'Risco Médio':
        return ('FFD966', '000000')
    if clf == 'Risco Alto':
        return ('C00000', 'FFFFFF')
    return ('BFBFBF', '000000')

def set_cell_text(cell, text, bold=False, size=9, color='000000', align='center'):
    cell.text = str(text)
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.runs[0]
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def get_credentials():
    info = st.secrets["gcp_service_account"]
    return Credentials.from_service_account_info(
        info,
        scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"]
    )

@st.cache_data(ttl=300)
def load_sheet(sheet_name, spreadsheet_id):
    credentials = get_credentials()
    gc = gspread.authorize(credentials)
    sh = gc.open_by_key(spreadsheet_id)
    ws = sh.worksheet(sheet_name)
    data = ws.get_all_records()
    return pd.DataFrame(data)

def prepare_scored_df(df_raw):
    df = df_raw.copy()
    for c in df.columns:
        if c in BASE_COLS:
            continue
        qnum = None
        nc = norm_text(c)
        for k, v in QUESTIONARIO_NORM.items():
            if nc == v:
                qnum = k
                break
        s = df[c].map(convert_answer_to_num)
        if qnum in NEGATIVAS:
            s = 6 - s
        df[c] = s
    return df

def get_company_subsets(df_raw, df_scored, company):
    company_norm = norm_text(company)
    mask = df_raw['NOME DA EMPRESA'].astype(str).map(norm_text) == company_norm
    sub_raw = df_raw[mask].copy()
    sub_scored = df_scored[mask].copy()
    if sub_raw.empty:
        raise ValueError(f'Empresa não encontrada: {company}')
    return sub_raw, sub_scored

def get_setores(sub_raw):
    if 'SETOR' not in sub_raw.columns:
        return ['GERAL']
    setores = sorted({str(x).strip() for x in sub_raw['SETOR'].dropna() if str(x).strip()})
    return setores if setores else ['GERAL']

def bloco_columns_present(df, inicio, fim):
    cols = []
    for i in range(inicio, fim + 1):
        nc = QUESTIONARIO_NORM[i]
        if nc in df.columns:
            cols.append(nc)
    return cols

def build_block_results(sub_scored):
    rows = []
    for bloco in BLOCOS:
        cols = bloco_columns_present(sub_scored, bloco['inicio'], bloco['fim'])
        if not cols:
            media = np.nan
        else:
            medias_por_respondente = sub_scored[cols].mean(axis=1, skipna=True)
            media = float(medias_por_respondente.mean()) if medias_por_respondente.notna().any() else np.nan
        rows.append({
            'Bloco': bloco['nome'],
            'Perguntas': f"{bloco['inicio']} a {bloco['fim']}",
            'Media': round(media, 2) if pd.notna(media) else np.nan,
            'Classificacao': classify_by_media(media)
        })
    return pd.DataFrame(rows)

def build_distribution_table(sub_raw):
    dist_rows = []
    total = len(sub_raw)
    for bloco in BLOCOS:
        first = True
        for i in range(bloco['inicio'], bloco['fim'] + 1):
            col = QUESTIONARIO_NORM[i]
            if col not in sub_raw.columns:
                continue
            serie = sub_raw[col].astype(str).map(norm_text)
            row = {'Bloco': bloco['nome'] if first else '', 'Numero': i, 'Pergunta': QUESTIONARIO[i]}
            first = False
            for txt, label in zip(ESCALA_TEXTOS, ESCALA_LABELS):
                n = int((serie == norm_text(txt)).sum())
                pct = round((n / total) * 100, 1) if total else 0
                row[label] = f'{n} ({pct}%)'
            dist_rows.append(row)
    return pd.DataFrame(dist_rows)

def get_meta(sub_raw):
    setores = ', '.join(sorted({str(x).strip() for x in sub_raw.get('SETOR', pd.Series(dtype=str)).dropna() if str(x).strip()}))
    cargos_col = 'CARGO ATUAL' if 'CARGO ATUAL' in sub_raw.columns else 'CARGO'
    cargos = ', '.join(sorted({str(x).strip() for x in sub_raw.get(cargos_col, pd.Series(dtype=str)).dropna() if str(x).strip()}))
    return {'Setores': setores or 'N/A', 'Cargos': cargos or 'N/A'}

def make_chart_blocks(res_blocos, company):
    dfp = res_blocos.copy()
    colors = [{'Risco Baixo': '#70AD47', 'Risco Médio': '#FFD966', 'Risco Alto': '#C00000'}.get(clf, '#BFBFBF') for clf in dfp['Classificacao']]
    plt.figure(figsize=(10, 5.5))
    bars = plt.bar(dfp['Bloco'], dfp['Media'].fillna(0), color=colors)
    plt.ylim(0, 5)
    plt.ylabel('Média')
    plt.title(f'{company} - Resultado por bloco')
    plt.xticks(rotation=10, ha='right')
    for bar, media, clf in zip(bars, dfp['Media'], dfp['Classificacao']):
        label = f'{media:.2f}\n{clf}' if pd.notna(media) else 'Sem dados'
        plt.text(bar.get_x() + bar.get_width() / 2, min((media if pd.notna(media) else 0) + 0.08, 4.9), label, ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    buffer = BytesIO()
    plt.savefig(buffer, format='png', dpi=220, bbox_inches='tight')
    buffer.seek(0)
    plt.close()
    return buffer

def add_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    style.font.size = Pt(10)

def doc_header(doc, company, sector):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'RELATÓRIO DE CONCLUSÃO DE RISCOS PSICOSSOCIAIS NO TRABALHO - {company} - {sector}')
    r.bold = True
    r.font.size = Pt(14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}')

def add_intro(doc, respondents, meta):
    doc.add_paragraph(f'Respostas analisadas: {respondents}')
    for k, v in meta.items():
        doc.add_paragraph(f'{k}: {v}')

def add_block_table(doc, res_blocos, display_sector):
    doc.add_heading('Análise dos fatores de risco psicossocial', level=1)
    doc.add_paragraph('Nesta seção, são detalhados os principais fatores identificados no questionário, com classificação visual por faixa de risco.')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['SETOR', 'SOBRECARGA DE TRABALHO', 'CONFLITO INTERPESSOAL E FALTA DE APOIO', 'RELAÇÕES FÍSICAS E EMOCIONAIS AO TRABALHO']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
    row_media = table.add_row().cells
    set_cell_text(row_media[0], display_sector, size=9)
    mapa_bloco = {r['Bloco']: r for _, r in res_blocos.iterrows()}
    ordem = ['Sobrecarga de trabalho', 'Conflito interpessoal e falta de apoio', 'Relações físicas e emocionais ao trabalho']
    for idx, nome in enumerate(ordem, start=1):
        media = mapa_bloco[nome]['Media']
        set_cell_text(row_media[idx], f"{media:.2f}" if pd.notna(media) else 'Sem dados', bold=True, size=10)
    row_risco = table.add_row().cells
    set_cell_text(row_risco[0], '', size=8)
    for idx, nome in enumerate(ordem, start=1):
        clf = mapa_bloco[nome]['Classificacao']
        fill, font_color = color_by_classificacao(clf)
        set_cell_text(row_risco[idx], clf.upper(), bold=True, size=8, color=font_color)
        set_cell_bg(row_risco[idx], fill)

def add_distribution_table(doc, dist_df):
    doc.add_heading('Distribuição de respostas por pergunta', level=1)
    doc.add_paragraph('Cada célula mostra a quantidade de respostas e a porcentagem correspondente em cada opção da escala.')
    cols = ['Bloco', 'Numero', 'Pergunta'] + ESCALA_LABELS
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(cols):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
    for _, row in dist_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            align = 'left' if col in ['Bloco', 'Pergunta'] else 'center'
            set_cell_text(cells[i], row.get(col, ''), size=8, align=align)

def add_recommendations(doc, res_blocos):
    doc.add_heading('Recomendações e ações corretivas/preventivas', level=1)
    for _, row in res_blocos.iterrows():
        bloco = row['Bloco']
        clf = row['Classificacao']
        if clf == 'Risco Alto':
            txt = f'{bloco}: recomenda-se ação imediata, análise das causas, validação com liderança e medidas corretivas formais.'
        elif clf == 'Risco Médio':
            txt = f'{bloco}: recomenda-se monitoramento, revisão de práticas internas e acompanhamento periódico.'
        elif clf == 'Risco Baixo':
            txt = f'{bloco}: manter acompanhamento contínuo e registrar boas práticas preventivas.'
        else:
            txt = f'{bloco}: revisar dados e preenchimento antes da conclusão final.'
        doc.add_paragraph(f'• {txt}')

def make_docx(company, sector, respondents, meta, res_blocos, dist_df, chart_buffer, display_sector):
    doc = Document()
    add_styles(doc)
    doc_header(doc, company, sector)
    add_intro(doc, respondents, meta)
    add_block_table(doc, res_blocos, display_sector)
    doc.add_paragraph('')
    doc.add_picture(chart_buffer, width=Inches(6.5))
    add_distribution_table(doc, dist_df)
    add_recommendations(doc, res_blocos)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_report(company, sector, df_raw, df_scored):
    sub_raw, sub_scored = get_company_subsets(df_raw, df_scored, company)
    if sector == 'ALL':
        setores = get_setores(sub_raw)
        if not setores:
            setores = ['GERAL']
    else:
        setores = [sector]
    buffers = []
    for setor in setores:
        if 'SETOR' in sub_raw.columns:
            mask = sub_raw['SETOR'].astype(str).map(norm_text) == norm_text(setor)
            setor_raw = sub_raw[mask].copy()
            setor_scored = sub_scored[mask].copy()
        else:
            setor_raw = sub_raw.copy()
            setor_scored = sub_scored.copy()
        if setor_raw.empty:
            continue
        respondents = len(setor_raw)
        meta = get_meta(setor_raw)
        display_sector = setor if setor else 'GERAL'
        res_blocos = build_block_results(setor_scored)
        dist_df = build_distribution_table(setor_raw)
        chart_buffer = make_chart_blocks(res_blocos, company)
        docx_buffer = make_docx(company, display_sector, respondents, meta, res_blocos, dist_df, chart_buffer, display_sector)
        buffers.append({'sector': display_sector, 'docx': docx_buffer, 'respondents': respondents})
    return buffers

def buscar_empresa_por_token(df_empresas, token):
    df_empresas = df_empresas.copy()
    df_empresas.columns = [norm_text(c) for c in df_empresas.columns]
    if 'TOKEN ÚNICO' not in df_empresas.columns or 'NOME' not in df_empresas.columns:
        return None
    linha = df_empresas[df_empresas['TOKEN ÚNICO'].astype(str).map(norm_text) == norm_text(token)]
    if linha.empty:
        return None
    return str(linha.iloc[0]['NOME']).strip()

st.title("📊 Gerador de Relatórios Psicossocial")

try:
    spreadsheet_id = st.secrets["google_forms"]["spreadsheet_id"]
except:
    st.error("Configure st.secrets com google_forms.spreadsheet_id e gcp_service_account.")
    st.stop()

params = st.query_params
token = params.get("token", None)
if isinstance(token, list):
    token = token[0]

if not token:
    st.error("🔒 Link inválido. Abra o painel usando o link enviado para sua empresa.")
    st.stop()

try:
    df_empresas = load_sheet("EMPRESAS", spreadsheet_id)
except Exception as e:
    st.error(f"Erro ao ler a aba EMPRESAS: {e}")
    st.stop()

empresa_selected = buscar_empresa_por_token(df_empresas, token)
if not empresa_selected:
    st.error("🔒 Token inválido ou empresa não encontrada.")
    st.stop()

st.success(f"✅ Empresa autenticada: {empresa_selected}")
st.markdown("---")

try:
    df_raw = load_sheet("Respostas", spreadsheet_id)
except Exception as e:
    st.error(f"Erro ao ler a aba de respostas: {e}")
    st.stop()

if df_raw.empty:
    st.warning("Nenhuma resposta encontrada na aba de respostas.")
    st.stop()

df_raw.columns = [norm_text(c) for c in df_raw.columns]

if 'NOME DA EMPRESA' not in df_raw.columns:
    st.error("A coluna 'NOME DA EMPRESA' não foi encontrada na aba de respostas.")
    st.stop()

df_scored = prepare_scored_df(df_raw)

company_norm = norm_text(empresa_selected)
mask_company = df_raw['NOME DA EMPRESA'].astype(str).map(norm_text) == company_norm
sub_raw_company = df_raw[mask_company].copy()
sub_scored_company = df_scored[mask_company].copy()

if sub_raw_company.empty:
    st.warning("Essa empresa ainda não possui respostas registradas.")
    st.stop()

setores = get_setores(sub_raw_company)
setores_display = ['ALL (Todos os setores)'] + setores
setor_selected = st.selectbox("Selecione o setor:", setores_display)
setor_value = setor_selected if setor_selected != 'ALL (Todos os setores)' else 'ALL'

st.metric("Respostas dessa empresa", len(sub_raw_company))

if st.button("📄 Gerar Relatório Word", type="primary", use_container_width=True):
    with st.spinner("⏳ Gerando relatório..."):
        try:
            buffers = generate_report(empresa_selected, setor_value, df_raw, df_scored)
            if not buffers:
                st.error("❌ Nenhuma resposta encontrada para essa combinação.")
            else:
                st.success(f"✅ Relatório gerado! {len(buffers)} setor(es).")
                for buf in buffers:
                    st.download_button(
                        label=f"⬇️ Download Word - {buf['sector']}",
                        data=buf['docx'].read(),
                        file_name=f"Relatorio_{empresa_selected}_{buf['sector']}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
        except Exception as e:
            st.error(f"❌ Erro ao gerar relatório: {e}")
